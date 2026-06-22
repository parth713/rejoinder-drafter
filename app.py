from __future__ import annotations

import os
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable

import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document


load_dotenv()

SUPPORTED_TYPES = ["pdf", "docx", "txt", "md"]
MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_TOTAL_CHARACTERS = 700_000

SYSTEM_PROMPT = """You are a senior arbitration counsel drafting a Rejoinder on behalf of the Claimant.

Your task is to produce a complete, filing-ready draft in Markdown based only on the supplied materials.

Mandatory drafting rules:
1. Use the Statement of Claim as the primary style and formatting template. Mirror its caption, party labels, defined terms, heading hierarchy, numbering style, tone, and prayer-for-relief structure wherever the source permits.
2. Answer the Statement of Defence comprehensively and point-by-point. Preserve its paragraph references when responding.
3. Reaffirm and incorporate the Claimant's positive case where appropriate, but do not merely repeat it.
4. Treat each supporting document as a potential Claimant exhibit. Use it only for propositions it actually supports and only if it is relevant to the Rejoinder.
5. Before assigning a new exhibit, check whether that document is already exhibited or clearly referenced as an exhibit in the Statement of Claim. Do not re-exhibit a document that is already exhibited; cite its existing exhibit reference instead.
6. For each relevant supporting document that is not already exhibited in the Statement of Claim, introduce and cite it as a new Claimant exhibit at the first paragraph that relies upon it. Identify the document by its date, parties/author, description, and uploaded filename where those details are available.
7. Follow the existing Claimant exhibit prefix and numbering scheme only when it is clear from the Statement of Claim, using the next unused sequential identifier. If the scheme or next number cannot be determined safely, use **[CLAIMANT TO ASSIGN EXHIBIT NUMBER: filename]**. Never guess an exhibit identifier.
8. Include a final **Schedule of New Exhibits to the Rejoinder** listing only the newly introduced exhibits, their proposed or placeholder identifiers, filenames, dates, descriptions, and the Rejoinder paragraphs in which they are first cited. Omit the schedule if no new exhibits are introduced.
9. Never invent facts, dates, quotations, contract clauses, exhibit numbers, procedural orders, admissions, evidence, legal authorities, or monetary figures.
10. If material information is missing, insert a conspicuous placeholder such as **[CLAIMANT TO CONFIRM: ...]** instead of guessing.
11. Distinguish allegations, denials, admissions, and matters requiring proof. Do not make an admission unless clearly supported by the supplied materials.
12. Where the Defence contains a contention that is not answered by the supplied materials, deny it where legally appropriate and state that the Respondent is put to strict proof; also add a confirmation placeholder if counsel input is required.
13. Output only the Rejoinder itself in valid Markdown. Do not include drafting commentary, a source summary, or a disclaimer.
14. Retain professional legal drafting and internal consistency throughout. The final document must be on behalf of the Claimant.
"""


@dataclass(frozen=True)
class ExtractedDocument:
    name: str
    role: str
    text: str


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"\n--- Page {index} ---\n{text}")
    extracted = "".join(pages).strip()
    if not extracted:
        raise ValueError("No selectable text was found. The PDF may be scanned and require OCR.")
    return extracted


def _extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text)
    for table_number, table in enumerate(document.tables, start=1):
        blocks.append(f"\n[Table {table_number}]")
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    extracted = "\n".join(blocks).strip()
    if not extracted:
        raise ValueError("No text was found in the DOCX file.")
    return extracted


def extract_uploaded_file(uploaded_file, role: str) -> ExtractedDocument:
    data = uploaded_file.getvalue()
    if len(data) > MAX_FILE_BYTES:
        raise ValueError("File exceeds the 20 MB limit.")

    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()
    if extension == "pdf":
        text = _extract_pdf(data)
    elif extension == "docx":
        text = _extract_docx(data)
    elif extension in {"txt", "md"}:
        try:
            text = data.decode("utf-8-sig").strip()
        except UnicodeDecodeError as exc:
            raise ValueError("Text files must use UTF-8 encoding.") from exc
        if not text:
            raise ValueError("The file is empty.")
    else:
        raise ValueError(f"Unsupported file type: .{extension}")

    return ExtractedDocument(name=uploaded_file.name, role=role, text=text)


def build_drafting_input(
    documents: list[ExtractedDocument], drafting_instructions: str
) -> str:
    document_sections = []
    for index, document in enumerate(documents, start=1):
        document_sections.append(
            f'<document index="{index}" role="{document.role}" '
            f'name="{document.name}">\n{document.text}\n</document>'
        )

    instructions = drafting_instructions.strip() or "No additional instructions supplied."
    joined_documents = "\n\n".join(document_sections)
    return f"""Draft the Claimant's complete Rejoinder using the materials below.

ADDITIONAL COUNSEL INSTRUCTIONS
{instructions}

SOURCE MATERIALS
{joined_documents}

Before producing the document, silently check that every substantive Defence contention has been addressed and that every factual assertion is grounded in the source materials. Then output only the final Markdown Rejoinder.
"""


def stream_rejoinder(
    client: OpenAI, model: str, drafting_input: str, max_output_tokens: int
) -> Iterable[str]:
    with client.responses.stream(
        model=model,
        instructions=SYSTEM_PROMPT,
        input=drafting_input,
        max_output_tokens=max_output_tokens,
    ) as stream:
        for event in stream:
            if event.type == "response.output_text.delta":
                yield event.delta


def _extract_all(claim_file, defence_file, supporting_files) -> list[ExtractedDocument]:
    documents = [
        extract_uploaded_file(claim_file, "Statement of Claim (primary format template)"),
        extract_uploaded_file(defence_file, "Statement of Defence (must be answered)"),
    ]
    documents.extend(
        extract_uploaded_file(
            file,
            "Supporting document (assess as a potential new Claimant exhibit)",
        )
        for file in supporting_files
    )
    total_characters = sum(len(document.text) for document in documents)
    if total_characters > MAX_TOTAL_CHARACTERS:
        raise ValueError(
            "The extracted materials are too large for a single drafting request. "
            "Please reduce them to the most relevant documents or split large files."
        )
    return documents


def main() -> None:
    st.set_page_config(page_title="Arbitration Rejoinder Drafter", page_icon="⚖️", layout="wide")

    st.title("Arbitration Rejoinder Drafter")
    st.caption(
        "Upload the pleadings and supporting record to prepare a Markdown draft on behalf of the Claimant."
    )

    with st.sidebar:
        st.header("Configuration")
        configured_model = os.getenv("OPENAI_MODEL", "").strip()
        if configured_model:
            st.success(f"Model: {configured_model}")
        else:
            st.error("OPENAI_MODEL is missing from .env")
        max_output_tokens = st.number_input(
            "Maximum output tokens",
            min_value=4_000,
            max_value=100_000,
            value=30_000,
            step=1_000,
            help="A higher limit permits a longer Rejoinder and may cost more.",
        )
        st.divider()
        st.warning(
            "AI-generated legal work requires review by qualified counsel. Uploaded text is sent to the configured model provider."
        )

    with st.form("rejoinder_form"):
        left, right = st.columns(2)
        with left:
            claim_file = st.file_uploader(
                "Statement of Claim *",
                type=SUPPORTED_TYPES,
                help="This document supplies the primary structure, numbering, and drafting style.",
            )
        with right:
            defence_file = st.file_uploader(
                "Statement of Defence *",
                type=SUPPORTED_TYPES,
                help="The draft will respond to this pleading point-by-point.",
            )

        supporting_files = st.file_uploader(
            "Supporting documents (optional)",
            type=SUPPORTED_TYPES,
            accept_multiple_files=True,
            help=(
                "Contracts, correspondence, procedural orders, witness material, or other useful records. "
                "Relevant documents not already exhibited in the Claim will be introduced as new exhibits."
            ),
        )
        drafting_instructions = st.text_area(
            "Additional drafting instructions (optional)",
            height=120,
            placeholder=(
                "Example: Use English law; preserve the existing exhibit references; "
                "focus on limitation and causation; do not revise the quantum section."
            ),
        )
        submitted = st.form_submit_button(
            "Draft Rejoinder", type="primary", use_container_width=True
        )

    if submitted:
        api_key = os.getenv("OPENAI_API_KEY", "").strip()
        model = os.getenv("OPENAI_MODEL", "").strip()
        if not claim_file or not defence_file:
            st.error("Please upload both the Statement of Claim and Statement of Defence.")
            return
        if not api_key or not model:
            st.error("Set OPENAI_API_KEY and OPENAI_MODEL in the .env file, then restart the app.")
            return

        try:
            with st.spinner("Reading the uploaded record…"):
                documents = _extract_all(claim_file, defence_file, supporting_files or [])
                drafting_input = build_drafting_input(documents, drafting_instructions)

            st.subheader("Draft Rejoinder")
            output_placeholder = st.empty()
            chunks: list[str] = []
            with st.status("Counsel AI is drafting…", expanded=False) as status:
                client = OpenAI(api_key=api_key)
                for chunk in stream_rejoinder(
                    client, model, drafting_input, int(max_output_tokens)
                ):
                    chunks.append(chunk)
                    output_placeholder.markdown("".join(chunks))
                status.update(label="Draft complete", state="complete")

            rejoinder = "".join(chunks).strip()
            if not rejoinder:
                raise RuntimeError("The model returned an empty draft.")
            st.session_state["rejoinder"] = rejoinder
        except Exception as exc:
            st.error(f"Could not draft the Rejoinder: {exc}")

    if st.session_state.get("rejoinder"):
        rejoinder = st.session_state["rejoinder"]
        with st.expander("Edit Markdown before downloading", expanded=False):
            edited = st.text_area("Markdown", value=rejoinder, height=500)
            if edited != rejoinder:
                st.session_state["rejoinder"] = edited
                rejoinder = edited
        st.download_button(
            "Download Rejoinder (.md)",
            data=rejoinder.encode("utf-8"),
            file_name="claimant_rejoinder.md",
            mime="text/markdown",
            use_container_width=True,
        )


if __name__ == "__main__":
    main()
